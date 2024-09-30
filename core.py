from docx import Document
import shutil
from datetime import datetime
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


choices = {
   "{name}": "Alae",
   "{id}": "12345",
   "{num boxes}": "3",
   "{flight}": "AT203",
   "tag1": "box1",
   "tag2": "box2"
}
document = "inputs/filldoc2.0.docx"


# Create a copy of the document
def docCopy(document):
   docCopy = "inputs/copy.docx"
   shutil.copy(document, docCopy)
   doc = Document(docCopy)
   return doc, docCopy


# Replace placeholders
def replace(document):
   for para in document.paragraphs:
       for run in para.runs:
           if "{date}" in run.text:
               print("Found placeholder: {date}")
               today = datetime.today()
               formatted_date = today.strftime("%B %d, %Y")
               run.text = run.text.replace("{date}", formatted_date)
           for placeholder in choices.keys():
               if placeholder in run.text:
                   print(f"Found placeholder: {run.text}")
                   run.text = run.text.replace(placeholder, choices[placeholder])


# Function to set row height to 0.32"
def set_row_height(row, height_in_inches):
   height_twips = int(height_in_inches * 1440)
   tr = row._tr
   trPr = tr.get_or_add_trPr()
   trHeight = OxmlElement('w:trHeight')
   trHeight.set(qn('w:val'), str(height_twips))
   trPr.append(trHeight)


# Add rows based on num_boxes
def add_table_rows(document):
   num_rows = int(choices["{num boxes}"])
   table = document.tables[0]
   for _ in range(num_rows):
       row = table.add_row()
       set_row_height(row, 0.32)


# Fill the table with placeholders and values
def fillTable(document):
   table = document.tables[0]
   filtered_choices = {k: v for k, v in choices.items() if k not in ["{name}", "{id}", "{num boxes}", "{flight}"]}
  
   while len(table.rows) - 1 < len(filtered_choices):
       table.add_row()


   for i, (placeholder, value) in enumerate(filtered_choices.items()):
       row = table.rows[i + 1]
       row.cells[0].text = placeholder
       row.cells[1].text = value
      
       for cell in row.cells:
           cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # Correct vertical alignment
           cell_paragraph = cell.paragraphs[0]
           cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Horizontal alignment
           run = cell_paragraph.runs[0] if cell_paragraph.runs else cell_paragraph.add_run()
           run.font.name = table.rows[0].cells[0].paragraphs[0].runs[0].font.name
           run.font.size = table.rows[0].cells[0].paragraphs[0].runs[0].font.size


# Main function
def main(document):
   doc, save_path = docCopy(document)
   replace(doc)
   add_table_rows(doc)
   fillTable(doc)
   doc.save(save_path)


main(document)

